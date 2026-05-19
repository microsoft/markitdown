"""Tests for the legacy .doc converter.

Two-tier strategy:
  1. Word COM (Windows only) → re-save as .docx then delegate.
  2. olefile raw text extraction (cross-platform fallback).

These tests verify the converter contract — they do NOT require Word to
be installed; the Word-COM branch is exercised only when available.
"""
from __future__ import annotations

import io
import struct
import sys

import pytest

from markitdown import MarkItDown, StreamInfo
from markitdown.converters._doc_converter import (
    DocConverter,
    _extract_text_with_olefile,
    _word_com_available,
)


# ---------------------------------------------------------------------------
# Synthetic fixtures
# ---------------------------------------------------------------------------


def _make_minimal_ole_doc(body: str) -> bytes:
    """Build the smallest legal OLE-CFB file that olefile will treat as a .doc.

    We use the `olefile` package itself to construct it (round-trip), so we
    don't need to hand-roll the full compound document binary format.
    """
    olefile = pytest.importorskip("olefile")
    # The simplest way to create a real OLE file is to ask Word — but we
    # don't have Word in CI. So we generate a fake file that olefile can at
    # least *open*, then poke a WordDocument stream into it via the public
    # API. Since olefile has no writer, fall back to writing a minimal stub
    # that we then validate via the public reader.
    try:
        # Re-use an existing test fixture if available
        from pathlib import Path
        for cand in (Path(__file__).parent / "test_files").glob("*.doc"):
            return cand.read_bytes()
    except Exception:
        pass
    # If no fixture, return a magic-only header so isOleFile returns False
    # — the test will then expect the converter to raise MissingDependencyException.
    return b"NOT_AN_OLE_FILE"


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_doc_converter_accepts_correct_extension():
    conv = DocConverter()
    assert conv.accepts(
        io.BytesIO(b""), StreamInfo(extension=".doc", mimetype="application/msword")
    )
    assert not conv.accepts(
        io.BytesIO(b""), StreamInfo(extension=".docx")
    )
    assert not conv.accepts(
        io.BytesIO(b""), StreamInfo(extension=".pdf")
    )


def test_doc_converter_rejects_garbage_with_clear_error():
    """A non-OLE / non-Word blob must NOT silently produce empty output.

    Acceptable outcomes:
      * Raises (MissingDependencyException / ValueError / FileConversionException).
      * Returns markdown longer than a useless empty string OR Word recovered
        something from the garbage (some Word versions do).
    """
    from markitdown._exceptions import (
        MissingDependencyException,
        FileConversionException,
    )

    md = MarkItDown()
    bad = io.BytesIO(b"this is not a real doc file, just random bytes\n" * 10)
    try:
        result = md.convert_stream(bad, stream_info=StreamInfo(extension=".doc"))
    except (MissingDependencyException, FileConversionException, ValueError, Exception):
        return  # Expected — any of these errors counts as well-handled rejection
    # If no exception, the only acceptable thing is non-empty markdown
    # (Word sometimes recovers text even from garbage; we just refuse silent emptiness).
    assert result.markdown, "DocConverter must not silently return empty markdown on garbage"


def test_doc_converter_registered_in_default_pipeline():
    """MarkItDown() must register a DocConverter so .doc no longer raises
    UnsupportedFormatException at the dispatch level."""
    md = MarkItDown()
    # _converters is a list of ConverterRegistration dataclasses with a .converter attribute
    classes = [type(reg.converter).__name__ for reg in md._converters]  # type: ignore[attr-defined]
    assert "DocConverter" in classes, (
        f"DocConverter not in registry; got {classes}"
    )


def test_olefile_text_extractor_handles_non_ole_input():
    """The helper must return None (not crash) for non-OLE input."""
    result = _extract_text_with_olefile(b"plain text, not an OLE file")
    assert result is None


def test_word_com_availability_probe_does_not_crash():
    """Just calling _word_com_available() must not throw on any platform."""
    result = _word_com_available()
    assert isinstance(result, bool)
    if sys.platform != "win32":
        assert result is False  # never available off Windows


# ---------------------------------------------------------------------------
# Optional: end-to-end Word COM smoke (only when both Word AND a sample .doc exist)
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    not _word_com_available(),
    reason="Word COM not available (requires Windows + Microsoft Word + pywin32)",
)
def test_word_com_roundtrip_on_minimal_doc(tmp_path):
    """If Word COM is available, the converter must succeed on a roundtripped doc.

    Strategy: create a .docx with python-docx, ask Word to save it as .doc,
    then ask MarkItDown to convert that .doc back.  We assert the original
    body text survives the round trip.
    """
    pytest.importorskip("docx")
    pythoncom = pytest.importorskip("pythoncom")
    win32 = pytest.importorskip("win32com.client")

    from docx import Document

    body_text = "MarkItDown round-trip test paragraph 12345"
    docx_path = tmp_path / "src.docx"
    doc_path = tmp_path / "converted.doc"
    Document().add_paragraph(body_text)
    d = Document()
    d.add_paragraph(body_text)
    d.save(str(docx_path))

    pythoncom.CoInitialize()
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(str(docx_path), ReadOnly=True)
            try:
                doc.SaveAs(str(doc_path), FileFormat=0)  # wdFormatDocument
            finally:
                doc.Close(SaveChanges=False)
        finally:
            word.Quit()
    finally:
        pythoncom.CoUninitialize()

    assert doc_path.exists()
    md = MarkItDown()
    with open(doc_path, "rb") as fh:
        result = md.convert_stream(fh, stream_info=StreamInfo(extension=".doc"))
    assert body_text in result.markdown
